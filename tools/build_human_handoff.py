#!/usr/bin/env python3
"""Build the human-facing Word handoff for the order-94 Ramsey result."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


PROJECT = Path("/Users/austingibbons/gather/maths/ramsey-template-94")
OUTPUT = PROJECT / "HANDOFF" / "R5_5_42078_Human_Handoff.docx"

SKILL_SCRIPTS = Path(
    "/Users/austingibbons/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "18324A"
INK = "20262E"
MUTED = "5E6873"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FB"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D8DEE6"
WHITE = "FFFFFF"
GREEN = "2E7D5B"
AMBER = "A66512"
RED = "A13B3B"

CONTENT_WIDTH_DXA = 9360

ORDER94_WORD = (
    "111212212222121112122121112122221221211131312321323233333333333333333333333333333232312321313"
)
ORDER453_WORD = (
    "122233332212111213121223131222132322323322213232232332211111221211122131322312122323132231213232"
    "312221323321111222332312212122333232211213212111223121122233232332332323322233233211111221111123"
    "131222213122313122121322231211222332221121322231212213132213122221313211111221111123323322233232"
    "332332323322211213221112123121122323332212122132332221111233231222132323121322313232212132231312"
    "21112122111112233232232312223323223231222131322121312111212233332221"
)

COLOR_1 = (
    "1 2 3 5 8 13 15 16 17 19 22 24 25 26 28 33 36 38 39 40 "
    "42 44 48 86 90 92"
)
COLOR_2 = (
    "4 6 7 9 10 11 12 14 18 20 21 23 27 29 30 31 32 34 35 37 "
    "45 47 50 52 82 84 87 89"
)
COLOR_3 = (
    "41 43 46 49 51 53 54 55 56 57 58 59 60 61 62 63 64 65 66 "
    "67 68 69 70 71 72 73 74 75 76 77 78 79 80 81 83 85 88 91 93"
)

assert len(ORDER94_WORD) == 93
assert len(ORDER453_WORD) == 452


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color: str) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def set_cell_border(cell, **edges: dict[str, str | int]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, color: str = MID_GREY, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_paragraph_border(
    paragraph,
    *,
    side: str,
    color: str,
    size: int = 8,
    space: int = 4,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name: str, size: float, color: str | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, "Calibri Light", 25, NAVY)
    title.font.bold = False
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 12, MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        set_style_font(style, "Calibri", size, color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Bullet 2"):
        style = styles[style_name]
        set_style_font(style, "Calibri", 11, INK)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.widow_control = True
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.50)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.75)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.25)

    caption = styles["Caption"]
    set_style_font(caption, "Calibri", 9, MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    set_style_font(code, "Courier New", 8.5, INK)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", 1)
    else:
        note = styles["Small Note"]
    set_style_font(note, "Calibri", 9.5, MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.05


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("RAMSEY TEMPLATE 94  |  HUMAN HANDOFF")
    set_run_font(run, "Calibri", 8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_paragraph_border(p, side="bottom", color=MID_GREY, size=4, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Independent review copy  •  ")
    set_run_font(run, "Calibri", 8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(p)
    for footer_run in p.runs:
        set_run_font(footer_run, "Calibri", 8.5)
        footer_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_rich_paragraph(
    doc: Document,
    parts: list[tuple[str, dict]],
    *,
    style: str | None = None,
    align=None,
    keep: bool = False,
):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_together = True
    for text, attrs in parts:
        run = p.add_run(text)
        if attrs.get("bold"):
            run.bold = True
        if attrs.get("italic"):
            run.italic = True
        if attrs.get("color"):
            run.font.color.rgb = RGBColor.from_string(attrs["color"])
        if attrs.get("font"):
            set_run_font(run, attrs["font"], attrs.get("size"))
        elif attrs.get("size"):
            run.font.size = Pt(attrs["size"])
    return p


def add_meta_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    label_run = p.add_run(label.upper() + "  ")
    set_run_font(label_run, "Calibri", 8.5)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(BLUE)
    value_run = p.add_run(value)
    set_run_font(value_run, "Calibri", 9.5)
    value_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_callout(doc: Document, kicker: str, headline: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    set_paragraph_shading(p, LIGHT_BLUE)
    set_paragraph_border(p, side="left", color=BLUE, size=20, space=8)
    run = p.add_run(kicker.upper())
    set_run_font(run, "Calibri", 9)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    set_paragraph_shading(p, LIGHT_BLUE)
    set_paragraph_border(p, side="left", color=BLUE, size=20, space=8)
    run = p.add_run(headline)
    set_run_font(run, "Calibri Light", 19)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, LIGHT_BLUE)
    set_paragraph_border(p, side="left", color=BLUE, size=20, space=8)
    run = p.add_run(body)
    set_run_font(run, "Calibri", 10.5)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_status_pill(paragraph, text: str, color: str) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, "Calibri", 9)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    weights: list[float],
    *,
    status_col: int | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        p = header_cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, "Calibri", 9.5)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_shading(header_cells[idx], DARK_BLUE)
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, "Calibri", 9.3)
            if idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            if status_col is not None and idx == status_col:
                run.bold = True
                value_lower = value.lower()
                if "complete" in value_lower or "confirmed" in value_lower:
                    run.font.color.rgb = RGBColor.from_string(GREEN)
                elif "pending" in value_lower or "required" in value_lower:
                    run.font.color.rgb = RGBColor.from_string(AMBER)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, PALE_BLUE)

    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_equation(doc: Document, text: str, size: float = 13) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, "Cambria Math", size)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def wrap_text(text: str, width: int = 74) -> list[str]:
    return [text[i : i + width] for i in range(0, len(text), width)]


def add_code_block(
    doc: Document,
    lines: list[str],
    *,
    caption: str | None = None,
    keep_together: bool = True,
) -> None:
    paragraphs = []
    for line in lines:
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.keep_together = keep_together
        set_paragraph_shading(p, LIGHT_GREY)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Courier New", 8.5)
        paragraphs.append(p)
    if paragraphs:
        paragraphs[0].paragraph_format.space_before = Pt(4)
        paragraphs[-1].paragraph_format.space_after = Pt(6)
        set_paragraph_border(paragraphs[0], side="top", color=MID_GREY, size=4, space=3)
        set_paragraph_border(paragraphs[-1], side="bottom", color=MID_GREY, size=4, space=3)
    if caption:
        doc.add_paragraph(caption, style="Caption")


def add_bullet(doc: Document, text: str, level: int = 1) -> None:
    style = "List Bullet" if level == 1 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_labelled_item(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(text)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    props = doc.core_properties
    props.title = "An Order-94 Rowley Template and the Bound R_5(5) >= 42,078"
    props.subject = "Human handoff for independent computational Ramsey review"
    props.author = "Prepared for independent review"
    props.keywords = "Ramsey numbers, multicolor Ramsey, Rowley template, exact certificate"
    props.comments = "Correctness, priority, and publication readiness are tracked separately."

    # First-page memo masthead.
    title = doc.add_paragraph(style="Title")
    title.add_run("An Order-94 Rowley Template\nand the Bound R₅(5) ≥ 42,078")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Human handoff — exact certificate, proof map, and review protocol")

    add_meta_line(doc, "Date", "23 July 2026")
    add_meta_line(doc, "Purpose", "Independent mathematical and computational review")
    add_meta_line(
        doc,
        "Status",
        "Finite certificate complete; priority and publication readiness pending external review",
    )
    add_meta_line(
        doc,
        "Companion repository",
        "/Users/austingibbons/gather/maths/ramsey-template-94",
    )
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(4)
    set_paragraph_border(rule, side="bottom", color=BLUE, size=10, space=3)

    add_callout(
        doc,
        "Lead result",
        "An explicit five-colouring of K₄₂₀₇₇ with no monochromatic K₅",
        "Consequently R₅(5) ≥ 42,078. The construction is explicit, cyclic, "
        "reconstructible from two compact distance words, and checked through "
        "independent exact implementations.",
    )

    doc.add_heading("Executive summary", level=1)
    add_rich_paragraph(
        doc,
        [
            (
                "We found an explicit effective Rowley (5,5,3) template of order 94 "
                "with extension parameter φ = 40. ",
                {},
            ),
            (
                "This is the single new finite object on which the result depends.",
                {"bold": True},
            ),
            (
                " Its two inherited colours remain K₅-free under period-93 repetition "
                "through the exact cutoff distance 368, while the third colour is a "
                "sum-free triangle-free template.",
                {},
            ),
        ],
    )
    add_rich_paragraph(
        doc,
        [
            (
                "Combining this word with an explicit cyclic three-colouring of K₄₅₃ "
                "recovered from Rowley’s archived 2022 ancillary data gives a cyclic "
                "five-colouring of K₄₂₀₇₇. Rowley’s published construction theorem "
                "then proves that no output colour contains K₅.",
                {},
            )
        ],
    )
    add_rich_paragraph(
        doc,
        [
            (
                "The checked survey benchmark is R₅(5) ≥ 41,626. The present construction "
                "would improve that recorded bound by 452. ",
                {},
            ),
            (
                "Correctness is high-confidence; priority is not yet claimed.",
                {"bold": True, "color": AMBER},
            ),
        ],
    )

    add_table(
        doc,
        ["Axis", "Current verdict", "Remaining human gate"],
        [
            ["Mathematics", "Complete exact chain", "Audit the source theorem specialization"],
            ["Finite certificate", "Confirmed locally", "Clean-room rerun on a second machine"],
            ["Current-record collision", "Not found in checked corpus", "Specialist priority search and author contact"],
            ["Publication readiness", "Near-ready", "External referee sign-off and attribution check"],
        ],
        [1.35, 1.55, 3.60],
        status_col=1,
    )

    add_page_break(doc)

    # Core mathematical statement and semantics.
    doc.add_heading("1. The mathematical statement", level=1)
    add_rich_paragraph(
        doc,
        [
            ("Theorem. ", {"bold": True, "color": DARK_BLUE}),
            (
                "There exists a five-colouring of the edges of K₄₂₀₇₇ with no "
                "monochromatic K₅. Therefore R₅(5) ≥ 42,078.",
                {"bold": True},
            ),
        ],
        keep=True,
    )
    add_equation(doc, "(94 − 1)(453 − 1) + 1 + 40 = 93·452 + 41 = 42,077")
    p = doc.add_paragraph()
    p.add_run(
        "The increment comes from extending Rowley’s previous order-93 template "
        "to order 94 while retaining φ = 40. In the compound construction, that "
        "one additional template position is amplified by the 452 nonzero distances "
        "of the second prototype."
    )

    doc.add_heading("2. Why the result matters", level=1)
    add_bullet(
        doc,
        "It supplies a concrete, fully explicit lower-bound construction rather than "
        "an asymptotic estimate or an uncertified search report.",
    )
    add_bullet(
        doc,
        "Relative to the checked revision-18 survey entry R₅(5) ≥ 41,626, it moves "
        "the bound to 42,078: a gain of 452 vertices.",
    )
    add_bullet(
        doc,
        "The witness is compact. A 93-symbol template word and a 452-symbol prototype "
        "word reconstruct all 42,076 distance colours of the final graph.",
    )
    add_bullet(
        doc,
        "Three distinct order-94 witnesses were found and independently accepted. "
        "Only the reflected canonical witness is needed for the theorem.",
    )

    doc.add_heading("3. Certificate semantics", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "For a word c:{1,…,m−1}→{1,2,3}, colour edge {i,j}, i<j, of Kₘ by c(j−i). "
        "Write A₁=c⁻¹(1), A₂=c⁻¹(2), and T=c⁻¹(3). The third class is a Rowley "
        "triangle-free template when m−1∈T and T has no positive triple a,b,a+b "
        "(including a=b). Its extension parameter is φ=min(T)−1."
    )
    p = doc.add_paragraph()
    p.add_run(
        "For inherited colours 1 and 2, period-(m−1) repetition assigns positive "
        "distance d the base colour"
    )
    add_equation(doc, "c(((d − 1) mod (m − 1)) + 1).")
    add_rich_paragraph(
        doc,
        [
            (
                "Load-bearing convention. ",
                {"bold": True, "color": RED},
            ),
            (
                "Residue zero is represented by base distance m−1, not by 0. "
                "For m=94, an inherited colour forbidding K₅ must be checked through "
                "distance 4(94−2)=368, meaning vertices 0,…,368: 369 vertices.",
                {},
            ),
        ],
    )

    add_page_break(doc)

    # Exact order-94 witness.
    doc.add_heading("4. The order-94 certificate", level=1)
    add_rich_paragraph(
        doc,
        [
            ("Canonical artifact. ", {"bold": True}),
            ("results/order94_t12.template", {"font": "Courier New", "size": 9}),
        ],
    )
    add_rich_paragraph(
        doc,
        [
            ("SHA-256. ", {"bold": True}),
            (
                "a3dd3415956277d68fc0197f6f5e35c38f2ed7936c66f11f658b0c57688abec1",
                {"font": "Courier New", "size": 8.7},
            ),
        ],
    )
    add_code_block(
        doc,
        wrap_text(ORDER94_WORD),
        caption="Figure 1. Canonical distance word c(1)c(2)…c(93); line breaks are typographic only.",
    )

    doc.add_heading("4.1 Distance classes", level=2)
    add_labelled_item(doc, "Colour 1.", COLOR_1)
    add_labelled_item(doc, "Colour 2.", COLOR_2)
    add_labelled_item(doc, "Template colour 3.", COLOR_3)

    doc.add_heading("4.2 Exact properties", level=2)
    add_table(
        doc,
        ["Property", "Exact value or check", "Evidence"],
        [
            ["Order and period", "m=94; period p=93", "Word length is exactly 93"],
            ["Template start", "min(T)=41; φ=40", "T∩{1,…,40}=∅"],
            ["Terminal distance", "93∈T", "Direct symbol check"],
            ["Triangle-free template", "No a,b,a+b∈T, including a=b", "Complete pair enumeration"],
            ["Inherited colour 1", "No K₅ on vertices 0,…,368", "Python and C++ exact search"],
            ["Inherited colour 2", "No K₅ on vertices 0,…,368", "Python and C++ exact search"],
            ["Reflection", "c(d)=c(41−d) and c(d)=c(134−d)", "Direct equality check"],
        ],
        [1.75, 2.75, 2.00],
    )
    p = doc.add_paragraph()
    p.add_run(
        "The two exact template verifiers also accept the same word through span 372, "
        "four complete periods. The theorem needs only span 368."
    )

    doc.add_heading("4.3 Robustness witnesses", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Two additional, distinct order-94 words are frozen as "
    )
    r = p.add_run("results/order94_direct.template")
    set_run_font(r, "Courier New", 9)
    p.add_run(" and ")
    r = p.add_run("results/order94_lazy.template")
    set_run_font(r, "Courier New", 9)
    p.add_run(
        ". Their template-class sizes are 35 and 37, compared with 39 for the "
        "canonical reflected word. They are corroborating evidence, not proof inputs."
    )

    add_page_break(doc)

    # Source theorem and exact compound.
    doc.add_heading("5. Source theorem and hypothesis map", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Fred Rowley’s Generalised Construction Theorem takes a linear "
        "U(k₁,…,kq−1,3;m) with a distinguished triangle-free template T and a "
        "second linear prototype of order n. It produces a linear Ramsey colouring "
        "of order"
    )
    add_equation(doc, "(m − 1)(n − 1) + 1 + φ,   where φ = min(T) − 1.")
    p = doc.add_paragraph()
    p.add_run(
        "For an inherited colour forbidding Kₖ, Rowley’s proof compresses any "
        "offending clique to maximum edge length at most (k−1)(m−2). We rely on "
        "that explicit per-colour proof bound, not on the theorem statement’s "
        "slightly different wording of its aggregate cutoff parameter."
    )

    add_table(
        doc,
        ["Theorem requirement", "Specialized discharge", "Bound artifact"],
        [
            ["First input", "Order-94 (5,5,3) linear word", "93-symbol canonical certificate"],
            ["Template property", "93∈T and T is sum-free", "Complete finite enumeration"],
            ["Extension parameter", "min(T)=41, hence φ=40", "Exact class list"],
            ["Inherited colours", "K₅-free through distance 368", "Independent Python/C++ checks"],
            ["Second input", "Linear cyclic (5,5,5;453) prototype", "452-symbol source extraction"],
            ["Second-input clique bound", "All three colours K₅-free", "Independent Python/C++ checks"],
        ],
        [2.00, 2.75, 1.75],
    )

    doc.add_heading("6. The exact compound colouring", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Let p=93. Let A₁,A₂,T be the classes of the order-94 word, and let "
        "B₁,B₂,B₃ be the classes of the order-453 prototype. For s∈{1,2}, define"
    )
    add_equation(
        doc,
        "A′ₛ = {ℓ+(μ−1)p : ℓ∈Aₛ, 1≤μ≤452} ∪ "
        "{ℓ+452p : ℓ∈Aₛ, 1≤ℓ≤40}.",
        11.5,
    )
    p = doc.add_paragraph()
    p.add_run(
        "For j∈{1,2,3}, relabel the prototype colours as output colours j+2 and put"
    )
    add_equation(
        doc,
        "B′ⱼ = {ℓ+(μ−1)p : ℓ∈T, μ∈Bⱼ}.",
        12,
    )
    p = doc.add_paragraph()
    p.add_run(
        "These five sets partition {1,…,42,076}. Colour edge {x,y}, x<y, of "
        "K₄₂₀₇₇ by the unique set containing y−x. This is the entire construction."
    )
    add_table(
        doc,
        ["Output colour", "Distance count", "Origin"],
        [
            ["1", "11,772", "First inherited template colour"],
            ["2", "12,676", "Second inherited template colour"],
            ["3", "5,304", "Prototype colour 1"],
            ["4", "7,566", "Prototype colour 2"],
            ["5", "4,758", "Prototype colour 3"],
            ["Total", "42,076", "Every positive distance exactly once"],
        ],
        [1.20, 1.45, 3.85],
    )
    add_rich_paragraph(
        doc,
        [
            ("Expanded word SHA-256. ", {"bold": True}),
            (
                "a97d5dce8a927db2889ba220119f9d4f3d1b88ee24d441f82a803a195ae8028d",
                {"font": "Courier New", "size": 8.7},
            ),
        ],
    )

    add_page_break(doc)

    # Full proof.
    doc.add_heading("7. Proof that the compound is K₅-free", level=1)
    doc.add_heading("7.1 Inherited colours", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Fix s∈{1,2} and suppose the compound contains a colour-s K₅. Translate "
        "its least vertex to 0 and write the remaining vertices as "
        "0<h₁<h₂<h₃<h₄. Every hᵢ and every hⱼ−hᵢ has a residue in Aₛ modulo p=93."
    )
    p = doc.add_paragraph()
    p.add_run(
        "If h₁>p, subtract p from all four positive coordinates. If a consecutive "
        "gap hᵢ₊₁−hᵢ exceeds p, subtract p from the entire upper tail "
        "hᵢ₊₁,…,h₄. Each operation preserves positivity, strict order, and every "
        "edge colour: cross-tail differences change by p, while other differences "
        "do not change."
    )
    p = doc.add_paragraph()
    p.add_run(
        "Iteration terminates with h₁≤p and every consecutive gap at most p. No "
        "inherited edge can have length divisible by p because residue p=93 lies "
        "in T. Hence h₁≤p−1 and every gap is at most p−1, so"
    )
    add_equation(doc, "h₄ ≤ 4(p−1) = 368.")
    p = doc.add_paragraph()
    p.add_run(
        "This would be a colour-s K₅ inside the exact span-368 graph rejected by "
        "both independent verifiers, a contradiction."
    )

    doc.add_heading("7.2 Prototype-derived colours", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Fix j∈{1,2,3} and suppose output colour j+2 contains a K₅. Translate its "
        "least vertex to 0. Every positive vertex has a unique representation"
    )
    add_equation(doc, "dᵢ = tᵢ + (μᵢ−1)p,   with tᵢ∈T and μᵢ∈Bⱼ.")
    p = doc.add_paragraph()
    p.add_run(
        "Take dᵢ<dₖ. If tᵢ<tₖ, the positive residue of dₖ−dᵢ is tₖ−tᵢ. It "
        "cannot lie in T, since tᵢ+(tₖ−tᵢ)=tₖ would violate sum-freeness. Thus an "
        "edge of the same prototype-derived colour forces tᵢ≥tₖ."
    )
    p = doc.add_paragraph()
    p.add_run(
        "If tᵢ=tₖ, the difference has prototype index μₖ−μᵢ. If tᵢ>tₖ, then"
    )
    add_equation(
        doc,
        "dₖ−dᵢ = (p−(tᵢ−tₖ)) + (μₖ−μᵢ−1)p,",
    )
    p = doc.add_paragraph()
    p.add_run(
        "and its prototype index is again μₖ−μᵢ. Therefore membership of every "
        "pairwise difference in the same output colour implies μₖ−μᵢ∈Bⱼ. The "
        "indices μᵢ are strictly increasing because the dᵢ increase while the tᵢ "
        "do not. Hence {0,μ₁,μ₂,μ₃,μ₄} is a monochromatic K₅ in the order-453 "
        "prototype, contradicting its exact verification."
    )
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    r = p.add_run("Conclusion. ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(
        "None of the five output colours contains K₅. Since the compound has order "
        "42,077, the definition of the diagonal five-colour Ramsey number gives "
        "R₅(5)≥42,078."
    )

    add_page_break(doc)

    # Verification.
    doc.add_heading("8. Verification and trust chain", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The search program is not part of the proof. The certificate is the frozen "
        "word plus exact verifiers. Positive objects were accepted only after "
        "algorithmically distinct implementations agreed."
    )

    add_table(
        doc,
        ["Object", "Checker A", "Checker B", "Result"],
        [
            ["Order-94 template", "Python integer bit masks", "C++ ordered candidate vectors", "VALID in both colours"],
            ["Order-453 prototype", "Python bit-mask K₅ search", "C++ vector-intersection search", "VALID in all three colours"],
            ["Compound word", "Block/residue generator", "Literal set-union reconstruction", "42,076 exact assignments"],
            ["Source extraction", "Hash-fixed ancillary XML parser", "Metadata and boundary-symbol checks", "Exact symbol agreement"],
        ],
        [1.60, 1.70, 1.80, 1.40],
    )

    doc.add_heading("8.1 Exact search evidence", level=2)
    add_table(
        doc,
        ["Finite object", "Colour 1 nodes", "Colour 2 nodes", "Colour 3 nodes"],
        [
            ["Order-94 template, Python", "550,261", "718,490", "—"],
            ["Order-453 prototype, Python", "1,040,265", "5,811,421", "782,585"],
            ["Order-453 prototype, C++", "237,855", "837,235", "201,199"],
        ],
        [2.40, 1.35, 1.35, 1.40],
    )
    p = doc.add_paragraph()
    p.add_run(
        "Different node counts are expected because the implementations use different "
        "candidate orderings and representations. Agreement of the final exhaustive "
        "verdict is the relevant invariant."
    )

    doc.add_heading("8.2 Regression coverage", level=2)
    add_bullet(
        doc,
        "Seven negative fixtures cover malformed length, missing terminal template "
        "distance, forbidden-prefix use, an additive template triple, insufficient "
        "span, and explicit repeated K₅ witnesses in both inherited colours.",
    )
    add_bullet(
        doc,
        "The canonical order-94 word is additionally overtested through span 372.",
    )
    add_bullet(
        doc,
        "The order-93 seed and order-453 prototype are re-extracted symbol-for-symbol "
        "from a SHA-256-fixed arXiv source archive.",
    )
    add_bullet(
        doc,
        "The expanded word is reconstructed independently from the displayed set "
        "unions; all distances 1,…,42,076 occur exactly once and the frozen hash matches.",
    )
    p = doc.add_paragraph()
    p.add_run(
        "Local rerun on 23 July 2026: the complete end-to-end suite exited successfully."
    )

    doc.add_heading("8.3 What was not used", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "A restricted 14-variable repair branch produced an UNSAT report and a DRAT "
        "file that has not been independently checked. That negative result is not "
        "used in the theorem. Likewise, no claim relies on search reproducibility: "
        "only the frozen witnesses and exact checkers matter."
    )

    add_page_break(doc)

    # Status and reviewer instructions.
    doc.add_heading("9. Honest status and known limitations", level=1)
    add_labelled_item(
        doc,
        "Correctness.",
        "The theorem chain is complete: both compact input objects are explicit, all "
        "finite hypotheses are exhaustively checked, the compound is exact, and the "
        "preservation proof is given above.",
    )
    add_labelled_item(
        doc,
        "Priority.",
        "Not established. Bounded searches found no order-94 collision, and the checked "
        "revision-18 survey still records 41,626. This supports target liveness only.",
    )
    add_labelled_item(
        doc,
        "Publication readiness.",
        "A short computational note is justified if a specialist confirms the literature "
        "status and a second machine reproduces the certificate.",
    )
    add_labelled_item(
        doc,
        "Direct-versus-compositional verification.",
        "The 42,077-vertex graph was not subjected to a naive all-vertex K₅ search. "
        "Its K₅-freeness follows from Rowley’s published theorem, whose hypotheses are "
        "discharged by exact finite checks. The full distance word was independently "
        "reconstructed.",
    )
    add_labelled_item(
        doc,
        "Implementation independence.",
        "The two clique checkers use different parsers and search representations but "
        "share the written semantics. A clean-room third implementation remains valuable.",
    )
    add_labelled_item(
        doc,
        "Machine independence.",
        "All checks reported here ran on one Apple Silicon machine. This is the final "
        "computational-trust gate, not an open proof obligation.",
    )

    doc.add_heading("10. Reviewer protocol", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Please approach the package adversarially. A useful review should try to "
        "invalidate the result in the following order."
    )
    add_labelled_item(
        doc,
        "Source theorem.",
        "Re-read arXiv:1912.01164v3, Definition 3.1 and Theorem 3.2 with proof "
        "(printed pp. 4–6). Audit the cutoff argument and every off-by-one convention.",
    )
    add_labelled_item(
        doc,
        "Prototype recovery.",
        "Re-extract ancillary column AB from the fixed 2022 source archive. Confirm that "
        "distances 1–452 are the prototype word and distance 453 begins template colour 4.",
    )
    add_labelled_item(
        doc,
        "Template semantics.",
        "Reimplement the checks without project imports. Include a=b in the sum-free "
        "test, represent residue zero by 93, inspect vertices 0,…,368, and test both "
        "inherited colours.",
    )
    add_labelled_item(
        doc,
        "Prototype clique search.",
        "Independently check all three colours of the order-453 word for K₅.",
    )
    add_labelled_item(
        doc,
        "Composition.",
        "Reconstruct the five distance sets from Section 6 and verify the order arithmetic, "
        "partition, and cyclic reflection.",
    )
    add_labelled_item(
        doc,
        "Priority.",
        "Search for unpublished, newly circulating, or differently named order-94 "
        "templates and for any post-survey improvement to R₅(5).",
    )
    add_rich_paragraph(
        doc,
        [
            (
                "Requested verdict format. ",
                {"bold": True, "color": DARK_BLUE},
            ),
            (
                "Report mathematical correctness, certificate trust, priority, and "
                "publication readiness as four separate axes. Lead with any counterexample "
                "or hypothesis mismatch.",
                {},
            ),
        ],
    )

    add_page_break(doc)

    # Appendix A: exact recovered prototype.
    doc.add_heading("Appendix A. Recovered order-453 prototype", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The following 452-symbol word is the complete cyclic three-colouring used as "
        "the second input. It was extracted from Rowley’s archived 2022 ancillary "
        "spreadsheet, sheet Paper_Sep_2022, column AB, distances 1–452."
    )
    add_rich_paragraph(
        doc,
        [
            ("Artifact. ", {"bold": True}),
            (
                "sources/rowley_exoo_order453.prototype",
                {"font": "Courier New", "size": 9},
            ),
        ],
    )
    add_rich_paragraph(
        doc,
        [
            ("SHA-256. ", {"bold": True}),
            (
                "19c97e6279c184f6f462786cadda4b7c9773d870a5b680a04eb1503ef384a2d0",
                {"font": "Courier New", "size": 8.7},
            ),
        ],
    )
    add_code_block(
        doc,
        wrap_text(ORDER453_WORD),
        caption="Figure A1. Complete order-453 prototype distance word; concatenate lines exactly.",
    )
    add_table(
        doc,
        ["Property", "Value", "Independent check"],
        [
            ["Order", "453", "Word length 452"],
            ["Colour counts", "136 / 194 / 122", "Direct symbol tally"],
            ["K₅ status", "Absent in all three colours", "Python and C++ exact searches"],
            ["Cyclic reflection", "v(d)=v(453−d)", "Direct equality check"],
            ["Source archive hash", "5041096…396b2", "Fixed arXiv source package"],
        ],
        [1.65, 2.40, 2.45],
    )
    p = doc.add_paragraph()
    p.add_run(
        "The historical attribution is not logically necessary once this explicit word "
        "is present and checked: the composition theorem needs only the finite prototype."
    )

    add_page_break(doc)

    # Appendix B: pseudocode and reproduction.
    doc.add_heading("Appendix B. Exact verifier pseudocode", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The core clique search is an exhaustive depth-first enumeration over strictly "
        "increasing candidate lists. Each clique appears exactly once."
    )
    pseudocode = [
        "Algorithm 1  Exact monochromatic clique search",
        "Input: ordered vertices V, adjacency predicate Adj, target size k",
        "Output: a k-clique, or NONE",
        "",
        "Search(prefix, candidates, need):",
        "    if need = 0:",
        "        return prefix",
        "    if |candidates| < need:",
        "        return NONE",
        "    while |candidates| >= need:",
        "        v ← least vertex in candidates",
        "        delete v from candidates",
        "        next ← {u in candidates : Adj(v,u)}",
        "        result ← Search(prefix followed by v, next, need−1)",
        "        if result is not NONE:",
        "            return result",
        "    return NONE",
        "",
        "return Search(empty sequence, V, k)",
    ]
    add_code_block(doc, pseudocode, keep_together=True)
    p = doc.add_paragraph()
    p.add_run(
        "Correctness invariant. At every recursive call, candidates are precisely the "
        "larger vertices adjacent to every vertex of the current prefix. The cardinality "
        "prune removes only branches that cannot attain size k. Therefore NONE is an "
        "exhaustive nonexistence certificate for the fixed graph supplied to the routine."
    )

    doc.add_heading("Appendix C. Reproduction command", level=1)
    add_code_block(
        doc,
        [
            "cd /Users/austingibbons/gather/maths/ramsey-template-94",
            "sh tests/run_end_to_end_checks.sh",
        ],
    )
    p = doc.add_paragraph()
    p.add_run(
        "Expected exit status: 0. The suite performs source extraction, both prototype "
        "checks, all three order-94 checks, the span-372 overtest, and independent "
        "compound reconstruction."
    )

    doc.add_heading("References", level=1)
    add_labelled_item(
        doc,
        "Rowley 2021.",
        "Fred Rowley, “A generalised linear Ramsey graph construction,” "
        "arXiv:1912.01164v3, Definition 3.1 and Theorem 3.2 with proof, printed pp. 4–6.",
    )
    add_labelled_item(
        doc,
        "Rowley 2022.",
        "Fred Rowley, “Improved Lower Bounds for Multicolour Ramsey Numbers using "
        "SAT-Solvers,” arXiv:2203.13476v3, Section 5 and Table 1, printed pp. 5–6, "
        "with ancillary sheet Paper_Sep_2022.",
    )
    add_labelled_item(
        doc,
        "Radziszowski 2026.",
        "Stanisław P. Radziszowski, “Small Ramsey Numbers,” Electronic Journal of "
        "Combinatorics Dynamic Survey DS1, revision 18, 24 April 2026, "
        "DOI 10.37236/21, Table XIa, printed p. 55.",
    )
    add_labelled_item(
        doc,
        "Campos–Pohoata 2026.",
        "Marcelo Campos and Cosmin Pohoata, “An update on multicolor Ramsey lower "
        "bounds,” arXiv:2601.15183v1. This asymptotic paper does not state the "
        "specific order-94 improvement.",
    )

    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.space_before = Pt(12)
    set_paragraph_border(p, side="top", color=MID_GREY, size=4, space=4)
    p.add_run(
        "End of handoff. The two compact input words, construction formula, proof, "
        "hashes, and review protocol above are sufficient to assess the result without "
        "trusting the search narrative."
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
